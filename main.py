from fastapi import FastAPI, Request, File, UploadFile
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
import os
import openai
import docx
from docx import Document
import uvicorn
from fastapi import HTTPException, Path
from fastapi import Depends

app = FastAPI()

openai.api_key = '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'

async def store_image_get_link(uploaded_file):
    file_location = os.getcwd() + "//Audio_file//" + uploaded_file.filename
    with open(file_location, "wb+") as file_object:
        file_object.write(uploaded_file.file.read())
    return file_location

def save_to_word(risk_recommendation):
    doc = Document()
    doc.add_heading('Generated Risk Recommendation', level=1)
    doc.add_paragraph(risk_recommendation)
    doc_path = os.getcwd() + "//Audio_file//" + 'risk_recommendation.docx'
    doc.save(doc_path)
    return doc_path

app.mount("/static", StaticFiles(directory="templates"), name="static")
templates = Jinja2Templates(directory="templates")

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("home.html", {"request": request})

@app.post("/upload/")
async def click_here(request: Request, file: UploadFile = File(...)):
    wave_path = await store_image_get_link(file)
    audio_file = open(wave_path, "rb")
    transcript = openai.Audio.transcribe(api_key=openai.api_key, model="whisper-1", file=audio_file, language="en")
    text_1 = transcript.text + "\n"
    print(text_1)
   
    prompt = (f"Given the provided transcript:\n\n{transcript}\n\n"
    "Based on industry-specific observations, your task is to generate a detailed report in the following format sequentially and not in tabular format: Observation in sequential manner, Hazard, Recommendation, IS code and it's respective clause governing the recommendation."
    "Using the transcript, you will get details like customer name, sum insured and industry type which needs to be shown in the top of the report.If the Industry type is not available, assign 'NA'. "
    "The detailed report that you will generate should contain a sequential mapping of the observations recorded with the hazards that the recommended observation could pose to the industry and respective measures required to combat the problem inferred from the observation and recommendations based on your knowledge. The recommendations should be backed by Indian Safety standard codes as per the industry and hazard type and mention these respective safety standard codes that needs to be adhered in the report as well."
    "Present this information row-wise in a professional report format"
    "Additionally, include point-wise risk recommendations aligned with environmental, health, industrial processes and safety norms related to that client specific industry type. "
    "Analyze the recommendations and ensure that the recommendations are based on relevant industry saftey standards code and provide them in the recommendation report."
    "Output the report containing the heading: Hazard and Recommendation in bold."
)
    response = openai.Completion.create(
        model="gpt-3.5-turbo-instruct",
        prompt=prompt,
        temperature=0.4,
        max_tokens=2000,
        n=1
    )

    risk_recommendation = response.choices[0].text.strip()
    print("Generated Risk Recommendation:")
    print(risk_recommendation)

    doc_path = save_to_word(risk_recommendation)
    return {"status": "200", "wave_path": wave_path, "doc_path": doc_path}

@app.get("/download/", response_class=HTMLResponse)
async def download(request: Request, doc_path: str):
    absolute_path = os.path.abspath(doc_path)
  
    return FileResponse(absolute_path, filename="risk_recommendation.docx")

if __name__=="__main__":
    uvicorn.run(app)
        